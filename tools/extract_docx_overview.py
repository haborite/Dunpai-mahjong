import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document


def text_of_runs(paragraph):
    return "".join(run.text for run in paragraph.runs).strip()


def table_text(table, max_rows=8):
    rows = []
    for row in table.rows[:max_rows]:
        rows.append([" ".join(cell.text.split()) for cell in row.cells])
    return rows


def detect_sections(paragraphs):
    sections = []
    refs_started = False
    for i, p in enumerate(paragraphs):
        txt = p["text"].strip()
        low = txt.lower()
        if not txt:
            continue
        if low in {"abstract", "introduction", "results", "discussion", "methods", "conclusion", "conclusions", "references", "acknowledgements", "acknowledgments"}:
            sections.append({"index": i, "heading": txt, "style": p["style"]})
        elif p["style"].lower().startswith("heading"):
            sections.append({"index": i, "heading": txt, "style": p["style"]})
        if low == "references":
            refs_started = True
    return sections, refs_started


def extract(path):
    doc = Document(str(path))
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = text_of_runs(p) or p.text.strip()
        if text:
            paragraphs.append({"index": i, "style": p.style.name if p.style else "", "text": text})

    tables = []
    for i, t in enumerate(doc.tables):
        tables.append({"index": i, "rows": len(t.rows), "cols": len(t.columns), "sample": table_text(t)})

    refs = []
    in_refs = False
    for p in paragraphs:
        low = p["text"].strip().lower()
        if low == "references":
            in_refs = True
            continue
        if in_refs:
            refs.append(p["text"])

    sections, has_refs = detect_sections(paragraphs)
    zip_flags = {}
    with zipfile.ZipFile(path) as z:
        names = set(z.namelist())
        zip_flags["comments"] = "word/comments.xml" in names
        zip_flags["footnotes"] = "word/footnotes.xml" in names
        zip_flags["endnotes"] = "word/endnotes.xml" in names
        document_xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        zip_flags["tracked_insertions"] = "<w:ins" in document_xml
        zip_flags["tracked_deletions"] = "<w:del" in document_xml
        zip_flags["figures_or_drawings"] = "<w:drawing" in document_xml or "<w:pict" in document_xml

    title_candidates = paragraphs[:20]
    word_count = sum(len(re.findall(r"\b[\w'-]+\b", p["text"])) for p in paragraphs)

    return {
        "path": str(path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "word_count_rough": word_count,
        "sections": sections[:80],
        "title_candidates": title_candidates,
        "first_200_paragraphs": paragraphs[:200],
        "tables": tables[:30],
        "references_count_rough": len(refs) if has_refs else 0,
        "references_sample": refs[:40],
        "zip_flags": zip_flags,
        "core_properties": {
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "author": doc.core_properties.author,
            "keywords": doc.core_properties.keywords,
            "last_modified_by": doc.core_properties.last_modified_by,
            "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
            "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
        },
    }


def main():
    if len(sys.argv) < 3:
        raise SystemExit("usage: extract_docx_overview.py OUT.json INPUT.docx [INPUT.docx ...]")
    out = Path(sys.argv[1])
    data = [extract(Path(p)) for p in sys.argv[2:]]
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(out)


if __name__ == "__main__":
    main()
